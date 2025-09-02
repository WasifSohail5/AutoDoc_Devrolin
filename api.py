from fastapi import FastAPI, HTTPException, BackgroundTasks, Request
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional
import os
import uuid
import logging
import traceback
import tempfile
from docx import Document
from docx2pdf import convert

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = FastAPI(title="Document Generator API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = r"E:\BSAI-5th\DataMining\devrolin\ETS\files"
OUTPUT_DIR = os.path.join(BASE_DIR, "generated_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)
class OfferLetterData(BaseModel):
    REF: str
    DATE: str
    NAME: str
    DURATION: str
    STARTDATE: str
    SUPNAME: str
    TASKS: str
    POSITION: str
    DEPARTMENT: str
    FROMANDTODATE: str
    TYPE: str
    RESPONSEDATE: str


class TerminationLetterData(BaseModel):
    REF: str
    DATE: str
    NAME: str
    POSITION: str
    TERMDATE: str
    LASTDAY: str


class CertificateData(BaseModel):
    NAME: str
    POSITION: str
    DURATION: str


# New models for experience letters
class AIMLExperienceLetterData(BaseModel):
    REF: str
    DATE: str
    NAME: str
    DURATION: str
    STARTDATE: str
    ENDDATE: str

class WebDevExperienceLetterData(BaseModel):
    REF: str
    DATE: str
    NAME: str
    DURATION: str
    STARTDATE: str
    ENDDATE: str

class GraphicDesignExperienceLetterData(BaseModel):
    REF: str
    DATE: str
    NAME: str
    DURATION: str
    STARTDATE: str
    ENDDATE: str

# Exception handler for detailed error messages
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    error_detail = f"Error: {str(exc)}\n{traceback.format_exc()}"
    logger.error(error_detail)
    return JSONResponse(
        status_code=500,
        content={"detail": str(exc), "trace": traceback.format_exc().split('\n')},
    )


def generate_document(template_name: str, data: dict, output_prefix: str):
    """Generate DOCX and PDF documents from template and data."""
    # Create unique filenames
    unique_id = str(uuid.uuid4())[:8]
    docx_filename = f"{output_prefix}_{unique_id}.docx"
    pdf_filename = f"{output_prefix}_{unique_id}.pdf"

    docx_path = os.path.join(OUTPUT_DIR, docx_filename)
    pdf_path = os.path.join(OUTPUT_DIR, pdf_filename)

    # Get template path
    template_path = os.path.join(TEMPLATES_DIR, template_name)

    # Log paths for debugging
    logger.info(f"Template path: {template_path}")
    logger.info(f"Output DOCX path: {docx_path}")
    logger.info(f"Output PDF path: {pdf_path}")

    # Check if template exists
    if not os.path.exists(template_path):
        raise HTTPException(status_code=500, detail=f"Template file not found: {template_name}")

    try:
        from docx import Document
        doc = Document(template_path)

        logger.info(f"Template loaded. Processing with data: {data}")

        # Enhanced placeholder replacement with additional logging
        for p in doc.paragraphs:
            paragraph_text = p.text

            # Log if paragraph contains any placeholders
            for key in data.keys():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph_text:
                    logger.info(f"Found placeholder {placeholder} in paragraph: {paragraph_text[:50]}...")

            # Replace placeholders
            inline = p.runs
            for i in range(len(inline)):
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in inline[i].text:
                        # Ensure value is a string and not None
                        str_value = str(value) if value is not None else ""
                        logger.info(f"Replacing {placeholder} with '{str_value}'")
                        inline[i].text = inline[i].text.replace(placeholder, str_value)

        # Also check for placeholders in tables if the document has any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in paragraph.text:
                                str_value = str(value) if value is not None else ""
                                logger.info(f"Replacing {placeholder} with '{str_value}' in table")
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str_value)

        # Save the new Word file
        logger.info(f"Saving Word document to {docx_path}...")
        doc.save(docx_path)

        # Convert Word file to PDF
        logger.info(f"Converting Word to PDF: {pdf_path}")
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
        except Exception as e:
            logger.error(f"PDF conversion failed: {str(e)}")
            # If PDF conversion fails, still return the DOCX file
            return {
                "docx_path": docx_path,
                "pdf_path": None,
                "docx_filename": docx_filename,
                "pdf_filename": None,
                "warning": "PDF conversion failed, only DOCX is available"
            }

        return {
            "docx_path": docx_path,
            "pdf_path": pdf_path,
            "docx_filename": docx_filename,
            "pdf_filename": pdf_filename
        }

    except Exception as e:
        logger.error(f"Document generation failed: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Document generation failed: {str(e)}")


# Certificate generation functions
def replace_placeholders(doc, replacements):
    for para in doc.paragraphs:
        text = para.text
        for key, value in replacements.items():
            if key in text:
                text = text.replace(key, value)
        if text != para.text:
            for run in para.runs:
                run.text = ""
            para.runs[0].text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    for key, value in replacements.items():
                        if key in text:
                            text = text.replace(key, value)
                    if text != para.text:
                        for run in para.runs:
                            run.text = ""
                        para.runs[0].text = text


def generate_certificate(template_path, output_pdf, name, position, duration):
    """Generate a certificate using the specified template and data."""
    doc = Document(template_path)
    replacements = {"{{NAME}}": name, "{{POSITION}}": position, "{{DURATION}}": duration}
    replace_placeholders(doc, replacements)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        temp_docx = tmp.name
    doc.save(temp_docx)

    convert(temp_docx, output_pdf)

    os.remove(temp_docx)

    logger.info(f"✅ Certificate generated at: {output_pdf}")
    return output_pdf


def cleanup_old_files(background_tasks: BackgroundTasks):
    """Add a task to delete old generated files."""

    def delete_old_files():
        try:
            files = [os.path.join(OUTPUT_DIR, f) for f in os.listdir(OUTPUT_DIR)]
            files.sort(key=os.path.getctime, reverse=True)

            for old_file in files[40:]:  # Keep 20 documents (both PDF and DOCX)
                if os.path.isfile(old_file):
                    os.remove(old_file)
        except Exception as e:
            logger.error(f"Cleanup failed: {str(e)}")

    background_tasks.add_task(delete_old_files)


@app.post("/generate/offer-letter")
async def generate_offer_letter(data: OfferLetterData, background_tasks: BackgroundTasks):
    logger.info(f"Generating offer letter for {data.NAME}")

    # Convert Pydantic model to dict with placeholders
    template_data = {
        "REF": str(data.REF) if data.REF else "",
        "DATE": str(data.DATE) if data.DATE else "",
        "NAME": str(data.NAME) if data.NAME else "",
        "DURATION": str(data.DURATION) if data.DURATION else "",
        "STARTDATE": str(data.STARTDATE) if data.STARTDATE else "",
        "SUPNAME": str(data.SUPNAME) if data.SUPNAME else "",
        "TASKS": str(data.TASKS) if data.TASKS else "",
        "POSITION": str(data.POSITION) if data.POSITION else "",
        "DEPARTMENT": str(data.DEPARTMENT) if data.DEPARTMENT else "",
        "FROMANDTODATE": str(data.FROMANDTODATE) if data.FROMANDTODATE else "",
        "TYPE": str(data.TYPE) if data.TYPE else "",
        "RESPONSEDATE": str(data.RESPONSEDATE) if data.RESPONSEDATE else ""
    }
    for key, value in template_data.items():
        logger.info(f"Template field {key}: '{value}'")

    try:
        result = generate_document("offer_template.docx", template_data, "offer_letter")
        cleanup_old_files(background_tasks)

        response = {
            "success": True,
            "message": "Offer letter generated successfully",
            "docx_url": f"/download/{result['docx_filename']}",
        }

        # Add PDF URL only if PDF was successfully generated
        if result.get("pdf_filename"):
            response["pdf_url"] = f"/download/{result['pdf_filename']}"
        if result.get("warning"):
            response["warning"] = result["warning"]

        return response

    except Exception as e:
        logger.error(f"Error generating offer letter: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/generate/termination-letter")
async def generate_termination_letter(data: TerminationLetterData, background_tasks: BackgroundTasks):
    logger.info(f"Generating termination letter for {data.NAME}")

    template_data = {
        "REFNO": str(data.REF) if data.REF else "",
        "DATE": str(data.DATE) if data.DATE else "",
        "NAME": data.NAME,
        "POSITION": data.POSITION,
        "TERMDATE": data.TERMDATE,
        "LASTDAY": data.LASTDAY
    }

    try:
        result = generate_document("Termination Letter.docx", template_data, "termination_letter")
        cleanup_old_files(background_tasks)

        response = {
            "success": True,
            "message": "Termination letter generated successfully",
            "docx_url": f"/download/{result['docx_filename']}",
        }

        if result.get("pdf_filename"):
            response["pdf_url"] = f"/download/{result['pdf_filename']}"
        if result.get("warning"):
            response["warning"] = result["warning"]

        return response

    except Exception as e:
        logger.error(f"Error generating termination letter: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/generate/certificate")
async def generate_certificate_endpoint(data: CertificateData, background_tasks: BackgroundTasks):
    logger.info(f"Generating certificate for {data.NAME}")

    try:
        # Create unique filenames for the certificate
        unique_id = str(uuid.uuid4())[:8]
        docx_filename = f"certificate_{unique_id}.docx"
        pdf_filename = f"certificate_{unique_id}.pdf"

        docx_path = os.path.join(OUTPUT_DIR, docx_filename)
        pdf_path = os.path.join(OUTPUT_DIR, pdf_filename)

        # Get template path
        template_path = os.path.join(TEMPLATES_DIR, "certificate.docx")

        # Generate the certificate using the dedicated function
        generate_certificate(
            template_path=template_path,
            output_pdf=pdf_path,
            name=data.NAME,
            position=data.POSITION,
            duration=data.DURATION
        )

        # Save a copy of the DOCX too (the function only returns the PDF)
        doc = Document(template_path)
        replacements = {
            "{{NAME}}": data.NAME,
            "{{POSITION}}": data.POSITION,
            "{{DURATION}}": data.DURATION
        }
        replace_placeholders(doc, replacements)
        doc.save(docx_path)

        cleanup_old_files(background_tasks)

        response = {
            "success": True,
            "message": "Experience certificate generated successfully",
            "docx_url": f"/download/{docx_filename}",
            "pdf_url": f"/download/{pdf_filename}"
        }

        return response

    except Exception as e:
        logger.error(f"Error generating certificate: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# New endpoints for experience letters
@app.post("/generate/experience-letter/aiml")
async def generate_aiml_experience_letter(data: AIMLExperienceLetterData, background_tasks: BackgroundTasks):
    logger.info(f"Generating AI/ML Experience letter for {data.NAME}")

    # Convert Pydantic model to dict with placeholders
    template_data = {
        "REF": str(data.REF) if data.REF else "",
        "DATE": str(data.DATE) if data.DATE else "",
        "NAME": str(data.NAME) if data.NAME else "",
        "DURATION": str(data.DURATION) if data.DURATION else "",
        "STARTDATE": str(data.STARTDATE) if data.STARTDATE else "",
        "ENDDATE": str(data.ENDDATE) if data.ENDDATE else "",
    }

    try:
        result = generate_document(r"E:\BSAI-5th\DataMining\devrolin\ETS\files\Experince_AI.docx", template_data, "aiml_experience_letter")
        cleanup_old_files(background_tasks)

        response = {
            "success": True,
            "message": "AI/ML Experience letter generated successfully",
            "docx_url": f"/download/{result['docx_filename']}",
        }

        # Add PDF URL only if PDF was successfully generated
        if result.get("pdf_filename"):
            response["pdf_url"] = f"/download/{result['pdf_filename']}"
        if result.get("warning"):
            response["warning"] = result["warning"]

        return response

    except Exception as e:
        logger.error(f"Error generating AI/ML Experience letter: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/generate/experience-letter/webdev")
async def generate_webdev_experience_letter(data: WebDevExperienceLetterData, background_tasks: BackgroundTasks):
    logger.info(f"Generating Web Development Experience letter for {data.NAME}")

    # Convert Pydantic model to dict with placeholders
    template_data = {
        "REF": str(data.REF) if data.REF else "",
        "DATE": str(data.DATE) if data.DATE else "",
        "NAME": str(data.NAME) if data.NAME else "",
        "DURATION": str(data.DURATION) if data.DURATION else "",
        "STARTDATE": str(data.STARTDATE) if data.STARTDATE else "",
        "ENDDATE": str(data.ENDDATE) if data.ENDDATE else "",
    }

    try:
        result = generate_document(r"E:\BSAI-5th\DataMining\devrolin\ETS\files\Experience_template.docx", template_data, "webdev_experience_letter")
        cleanup_old_files(background_tasks)

        response = {
            "success": True,
            "message": "Web Development Experience letter generated successfully",
            "docx_url": f"/download/{result['docx_filename']}",
        }

        # Add PDF URL only if PDF was successfully generated
        if result.get("pdf_filename"):
            response["pdf_url"] = f"/download/{result['pdf_filename']}"
        if result.get("warning"):
            response["warning"] = result["warning"]

        return response

    except Exception as e:
        logger.error(f"Error generating Web Development Experience letter: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/generate/experience-letter/graphic-design")
async def generate_graphic_design_experience_letter(data: GraphicDesignExperienceLetterData,
                                                    background_tasks: BackgroundTasks):
    logger.info(f"Generating Graphic Design Experience letter for {data.NAME}")

    # Convert Pydantic model to dict with placeholders
    template_data = {
        "REF": str(data.REF) if data.REF else "",
        "DATE": str(data.DATE) if data.DATE else "",
        "NAME": str(data.NAME) if data.NAME else "",
        "DURATION": str(data.DURATION) if data.DURATION else "",
        "STARTDATE": str(data.STARTDATE) if data.STARTDATE else "",
        "ENDDATE": str(data.ENDDATE) if data.ENDDATE else "",
    }

    try:
        result = generate_document(r"E:\BSAI-5th\DataMining\devrolin\ETS\files\Experience_Graphic.docx", template_data,
                                   "graphic_design_experience_letter")
        cleanup_old_files(background_tasks)

        response = {
            "success": True,
            "message": "Graphic Design Experience letter generated successfully",
            "docx_url": f"/download/{result['docx_filename']}",
        }

        # Add PDF URL only if PDF was successfully generated
        if result.get("pdf_filename"):
            response["pdf_url"] = f"/download/{result['pdf_filename']}"
        if result.get("warning"):
            response["warning"] = result["warning"]

        return response

    except Exception as e:
        logger.error(f"Error generating Graphic Design Experience letter: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join(OUTPUT_DIR, filename)

    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        raise HTTPException(status_code=404, detail="File not found")

    # Determine content type based on file extension
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if filename.endswith(
        ".docx") else "application/pdf"

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type=content_type
    )


@app.get("/")
async def root():
    return {"message": "Document Generator API is running. Access the API documentation at /docs"}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)